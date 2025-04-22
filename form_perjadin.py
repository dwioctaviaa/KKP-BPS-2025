import os
import io
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.enum.style import WD_STYLE_TYPE
import mysql.connector
from flask import Blueprint, render_template

app = Flask(__name__)

# Folder penyimpanan template dan upload gambar
TEMPLATE_FOLDER = './templates'
UPLOAD_FOLDER = './static/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

perjadin_routes = Blueprint('perjadin_routes', __name__)

@perjadin_routes.route('/form_perjadin')
def form_laporan():
    return render_template('form_lapangan.html')



# Pastikan folder upload ada
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def simpan_ke_database(nama_file, jenis_form, nama_user, path_file=None):
    db = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="db_perjadin"
    )
    cursor = db.cursor()
    query = """
        INSERT INTO history_file (nama_file, jenis_form, tanggal_generate, nama_user, path_file)
        VALUES (%s, %s, %s, %s, %s)
    """
    data = (nama_file, jenis_form, datetime.now(), nama_user, path_file)
    cursor.execute(query, data)
    db.commit()
    cursor.close()
    db.close()
    return render_template('dashboard.html', history=data)

@perjadin_routes.route('/dashboard')
def dashboard():
    db = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="db_perjadin"
    )
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM history_file ORDER BY tanggal_generate DESC")
    history = cursor.fetchall()
    cursor.close()
    db.close()
    return render_template('dashboard.html', history=history)


def format_waktu(waktu):
    """Mengubah waktu dari 'HH:MM' menjadi 'HH.MM WITA'."""
    if not waktu:
        return "Waktu tidak tersedia"
    
    try:
        waktu_obj = datetime.strptime(waktu, "%H:%M")  # Konversi input HTML ke datetime
        return waktu_obj.strftime("%H.%M") + " WITA"  # Ubah format ke "HH.MM WITA"
    except ValueError:
        return "Format waktu salah"

# Fungsi untuk mengubah format tanggal menjadi '12 Maret 2023'
def format_tanggal(tanggal_str):
    bulan_inggris_ke_indonesia = {
        "January": "Januari", "February": "Februari", "March": "Maret",
        "April": "April", "May": "Mei", "June": "Juni",
        "July": "Juli", "August": "Agustus", "September": "September",
        "October": "Oktober", "November": "November", "December": "Desember"
    }
    try:
        dt = datetime.strptime(tanggal_str, '%Y-%m-%d')  # Pastikan format input tanggal sesuai 'YYYY-MM-DD'
        tanggal_format = dt.strftime('%d %B %Y')
        for eng, ind in bulan_inggris_ke_indonesia.items():
            tanggal_format = tanggal_format.replace(eng, ind)
        return tanggal_format
    except ValueError:
        return "Format tanggal salah, pastikan format 'YYYY-MM-DD' digunakan"


# Fungsi untuk mengubah font seluruh dokumen ke Arial 11
def change_font(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Fungsi untuk menghapus baris kosong dalam tabel
def remove_empty_rows(doc):
    for table in doc.tables:
        rows_to_remove = [row for row in table.rows if all(cell.text.strip() == "" for cell in row.cells)]
        for row in rows_to_remove:
            table._tbl.remove(row._tr)

# Fungsi untuk membuat tabel tanda tangan
def create_signature_table(doc, data):
    bagian_ketua_tim = data["bagian_ketua_tim"]
    ketua_tim = data["ketua_tim"]
    nama_petugas = data["nama_petugas"]
    lokasi = data["lokasi"]
    tanggal_surat=data["tanggal_surat"]
    nip_ketua = data["nip_ketua"]
    nip_petugas = data["nip_petugas"]
    
    # Kode untuk membuat tabel tanda tangan
    table_paragraph = doc.add_paragraph()
    table_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=2, cols=2)
    table_paragraph._p.addnext(table._tbl)

    # Baris pertama (judul kolom)
    table.cell(0, 0).text = f"Mengetahui,\n{bagian_ketua_tim}\n\n\n"
    table.cell(0, 1).text = f"{lokasi}, {format_tanggal(tanggal_surat)}\nPetugas,"

    # Baris kedua (nama Ketua Tim dan Petugas)
    format_existing_text(table.cell(1, 0), ketua_tim)  # Ketua Tim
    format_existing_text(table.cell(1, 1), nama_petugas)  # Nama Petugas
    # Menambahkan NIP
    table.cell(1, 0).add_paragraph(f"NIP. {nip_ketua}")
    table.cell(1, 1).add_paragraph(f"NIP. {nip_petugas}")
    # Format teks dan perataan
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.font.size = Pt(10)

    tbl = table._tbl
    for row in tbl.iter_tcs():
        tc_pr = row.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            borders.append(border)
        tc_pr.append(borders)
def format_existing_text(cell, text):
    """Mengubah teks yang sudah ada di dalam tabel menjadi bold & underline."""
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()  # Pastikan ada paragraf
    else:
        paragraph = cell.paragraphs[0]  # Ambil paragraf pertama
    
    paragraph.clear()  # Hapus teks lama agar hanya teks yang diinginkan yang diformat
    run = paragraph.add_run(text)  # Tambahkan teks yang sudah ada
    run.bold = True
    run.underline = True
def create_signature_table_generate_word(data_list, doc): 
    # Membuat tabel dengan 3 kolom
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Menambahkan header ke tabel
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Jam'
    hdr_cells[1].text = 'Kegiatan'
    hdr_cells[2].text = 'Jam Kunjungan'

    # Menambahkan data ke dalam tabel
    for item in data_list:  
        row = table.add_row()  # Tambahkan baris baru dalam tabel
        row.cells[0].text = f"{item['waktu_awal']} - {item['waktu_akhir']}"
        row.cells[1].text = item['kegiatan']  # Kolom Kegiatan
        row.cells[2].text = item['jam_kunjungan']  # Kolom Jam Kunjungan

    return doc

def create_jadwal_perjalanan_dinas(doc, data):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Jadwal Kegiatan")
    run.bold = True  # Untuk teks tebal
    # Mengatur agar teks berada di tengah
    paragraph.alignment = 1
        # Membuat tabel dengan 2 kolom
    table = doc.add_table(rows=9, cols=2)
    # Menghilangkan border tabel
    tbl = table._element
    for cell in tbl.xpath(".//w:tc"):
        tc_pr = cell.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')  # Menggunakan namespace yang benar
            tc_borders.append(border)
        tc_pr.append(tc_borders)

    # Menambahkan header tabel
    table.cell(0, 0).text = 'Surat Perjalanan Dinas' 
    table.cell(1, 0).text = 'Nomor                                                 :'
    table.cell(2, 0).text = 'Tanggal                                               :'
    table.cell(3, 0).text = 'Nama Pelaksana Perjalanan Dinas    :'
    table.cell(4, 0).text = 'NIP                                                      :'
    table.cell(5, 0).text = 'Pangkat / Golongan                            :'
    table.cell(6, 0).text = 'Jabatan / Instansi                                :'
    table.cell(7, 0).text = 'Maksud Perjalanan Dinas                   :'
    table.cell(8, 0).text = 'Tanggal Pelaksanaan Perjalanan Dinas:'
                                                

    # Memasukkan data ke dalam kolom kedua
    table.cell(1, 1).text = data.get('nomor_surat', '')
    tanggal_surat = data.get('tanggal_surat', '')  # Ambil nilai dari data
    if tanggal_surat:  # Jika tidak kosong, format tanggalnya
        tanggal_surat = format_tanggal(tanggal_surat)

    table.cell(2, 1).text = tanggal_surat  # Masukkan ke dalam tabel
    table.cell(3, 1).text = data.get('nama_petugas', '')
    table.cell(4, 1).text = data.get('nip_petugas', '')  # Jika Anda menambahkan data 'nip'
    table.cell(5, 1).text = data.get('pangkat', '')
    table.cell(6, 1).text = data.get('jabatan_petugas', '')  # Anda bisa menambahkan data 'jabatan_instansi'
    table.cell(7, 1).text = data.get('maksud_perjalanan_dinas', '')
    table.cell(8, 1).text = data.get('tanggal', '')
    # Menyesuaikan tampilan tabel agar rapi dan memberikan jarak antar tabel
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Ratakan ke kiri
                paragraph.space_after = Pt(10)  # Menambahkan jarak antar tabel

    # Menambahkan jarak setelah tabel
    doc.add_paragraph("\n")  

    return doc
@app.route('/')
def home():
    return render_template("index.html")
formatted_waktu = None
@app.route('/generate', methods=['POST'])
def generate():
    # Ambil data dari form
    nama_petugas = request.form.get('nama_petugas', '')
    nip_petugas = request.form.get('nip_petugas', '')
    jabatan_petugas = request.form.get('jabatan_petugas', '')
    nip_ketua= request.form.get('nip_ketua', '')
    dasar_perlaksanaan = request.form.get('dasar_perlaksanaan', '')
    nomor_surat=request.form.get('nomor_surat', '')
    waktu = request.form.get("waktu")  # Ambil input dari form
    formatted_waktu = format_waktu(waktu)  # Ubah format waktu
    lokasi = request.form.get('lokasi')
    personil = request.form.get('personil', '')
    pokok_hasil = request.form.get('pokok_hasil', '')
    tanggal = request.form.get('tanggal')
    if not tanggal:
        return "Tanggal tidak boleh kosong", 400  # Menangani jika tanggal kosong
    
    # Format tanggal
    formatted_tanggal = format_tanggal(tanggal)
    bagian_ketua_tim = request.form.get('bagian_ketua_tim', '')
    ketua_tim= request.form.get('ketua_tim', '')
    template_path = os.path.join(TEMPLATE_FOLDER, 'template_perjadin.docx')
    nomor_surat=request.form.get('nomor_surat', '')
    pangkat=request.form.get('pangkat', '')
    maksud_perjalanan_dinas =request.form.get('maksud_perjalanan_dinas', '')
    tanggal_surat= request.form.get('tanggal_surat', '')
    waktu_awal = request.form.getlist('waktu_awal[]')
    waktu_akhir = request.form.getlist('waktu_akhir[]')
    kegiatan = request.form.getlist('kegiatan[]')
    jam_kunjungan = request.form.getlist('jam_kunjungan[]')
    add_row = request.form.get('add_row')  # Tambahkan parameter untuk menambah baris baru

    # Menyimpan data form
    data = {
        "bagian_ketua_tim": bagian_ketua_tim,
        "ketua_tim": ketua_tim,
        "dasar_perlaksanaan": dasar_perlaksanaan,
        "nama_petugas": nama_petugas,
        "lokasi": lokasi,
        "tanggal": formatted_tanggal,
        "nip_ketua": nip_ketua,
        "nip_petugas": nip_petugas,
        "pokok_hasil": pokok_hasil,
        "waktu": waktu,
        "nomor_surat": nomor_surat,
        "nip_petugas": nip_petugas,
        "pangkat": pangkat,
        "jabatan_petugas":jabatan_petugas,
        "maksud_perjalanan_dinas":maksud_perjalanan_dinas,
        "tanggal_surat":tanggal_surat,
        "formatted_tanggal":formatted_tanggal,
        "add_row":add_row
    }

    data_list = [
            {"waktu_awal": waktu_awal[i], "waktu_akhir": waktu_akhir[i], "kegiatan": kegiatan[i], "jam_kunjungan": jam_kunjungan[i]}
            for i in range(len(waktu_awal))
        ]

    if not os.path.exists(template_path):
        return "Template tidak ditemukan."
    doc = Document(template_path)
      # Masukkan data ke dalam tabel Word untuk Bagian A
    for table in doc.tables:
        for row in table.rows:
            if 'Nama' in row.cells[0].text:
                row.cells[1].text = nama_petugas
            elif 'NIP' in row.cells[0].text:
                row.cells[1].text = nip_petugas
            elif 'Jabatan' in row.cells[0].text:
                row.cells[1].text = jabatan_petugas

    # Masukkan data ke dalam tabel pada baris kedua
    for table in doc.tables:
        if len(table.rows) >= 2:  # Pastikan tabel memiliki minimal dua baris
            title_cell = table.rows[0].cells[0].text.strip()
            if title_cell in "B. Dasar Perlaksanaan:":
                cell = table.rows[1].cells[0]
                cell.text = data.get('dasar_perlaksanaan', '')
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    paragraph.paragraph_format.first_line_indent = Inches(0.3)
                    paragraph.paragraph_format.left_indent = Inches(0.2)  # Indentasi untuk semua teks 
    for table in doc.tables:
        for row in table.rows:
            # Memeriksa apakah kolom pertama mengandung 'C. Waktu'
            if 'Waktu' in row.cells[0].text:
                row.cells[1].text = f"Perjalanan Dinas dilakukan pada tanggal {formatted_tanggal}."
            # Memeriksa apakah kolom pertama mengandung 'D. Lokasi'
            elif 'Lokasi' in row.cells[0].text:
                row.cells[1].text = lokasi
            # Memeriksa apakah kolom pertama mengandung 'E. Personil'
            elif 'Personil' in row.cells[0].text:
                row.cells[1].text = personil
    for table in doc.tables:
        if len(table.rows) >= 2:  # Pastikan tabel memiliki minimal dua baris
            title_cell = table.rows[0].cells[0].text.strip()
            if title_cell in "F. Pokok-pokok Hasil Kegiatan:":
                cell = table.rows[1].cells[0]
                cell.text = pokok_hasil
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    paragraph.paragraph_format.first_line_indent = Inches(0.3)
                    paragraph.paragraph_format.left_indent = Inches(0.2) 

    # Simpan gambar yang diunggah
    for para in doc.paragraphs:
        if para.text.strip() == "Dokumentasi":
            # Mengambil semua gambar yang diunggah
            dokumentasi_files = request.files.getlist('dokumentasi')  # Mendapatkan daftar gambar yang diunggah
            
            if dokumentasi_files:
                for dokumentasi in dokumentasi_files:
                    if dokumentasi and dokumentasi.filename != '':
                        # Membaca gambar sebagai byte stream
                        gambar_bytes = dokumentasi.read()
                        
                        # Mengonversi byte stream ke objek file-like
                        image_stream = io.BytesIO(gambar_bytes)
                        
                        # Menambahkan gambar ke dalam dokumen setelah judul
                        paragraph = doc.add_paragraph()  # Membuat paragraf kosong untuk gambar
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Ratakan gambar ke tengah
                        run = paragraph.add_run()
                        
                        try:
                            # Menambahkan gambar dengan memastikan format yang tepat
                            run.add_picture(image_stream, width=Inches(5), height=Inches(4))
                        except Exception as e:
                            # Menangani kesalahan jika gambar tidak bisa dimasukkan
                            print(f"Error adding image: {e}")
                        
                doc.add_page_break()
            break
    doc = create_jadwal_perjalanan_dinas(doc, data)
    # Menambahkan tabel Jam, kegiatan dan Tanda Tangan
    doc = create_signature_table_generate_word(data_list, doc)
    
    create_signature_table(doc, data)  # Menambahkan tabel tanda tangan
    # Format dokumen

    change_font(doc)
    remove_empty_rows(doc)

    # Simpan dokumen ke dalam memori
    word_mem = io.BytesIO()
    doc.save(word_mem)
    word_mem.seek(0)

    filename = f"Laporan_{nama_petugas}.docx"
    simpan_ke_database(
        nama_file=filename,
        jenis_form="perjadin",
        nama_user=nama_petugas,
        path_file=None  # atau bisa isi "Dikirim dari memori"
    )


    return send_file(word_mem, as_attachment=True, download_name=f"Laporan_{nama_petugas}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True)
