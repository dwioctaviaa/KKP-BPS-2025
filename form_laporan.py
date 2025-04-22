from flask import Flask, render_template, request, send_file, session, redirect, url_for, after_this_request
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
import comtypes.client
from datetime import datetime
from comtypes import CoInitialize
from time import sleep
import mysql.connector
from flask import Blueprint, render_template

app = Flask(__name__)
app.secret_key = "secret_key"  
TEMPLATE_FOLDER = 'templates'
ketua_tim_list = []

laporan_routes = Blueprint('laporan_routes', __name__)

@laporan_routes.route('/form_laporan', methods=['GET', 'POST'])
def form_laporan():
    if request.method == 'POST':
        # proses data form di sini
        pass
    return render_template('form_lapangan.html')


def simpan_ke_database(nama_file, jenis_form, nama_user, path_file=None):
    db = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="db_perjadin"
    )
    cursor = db.cursor()
    query = """
        INSERT INTO history_file (nama_file, jenis_form, tanggal_generate, nama_user, path_file)
        VALUES (%s, %s, %s, %s, %s)
    """
    data = (nama_file, jenis_form, datetime.now(), nama_user, path_file)
    cursor.execute(query, data)
    db.commit()
    cursor.close()
    db.close()

@laporan_routes.route('/dashboard')
def dashboard():
    db = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="db_perjadin"
    )
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM history_file ORDER BY tanggal_generate DESC")
    history = cursor.fetchall()
    cursor.close()
    db.close()
    return render_template('dashboard.html', history=history)


def change_font(doc):
    """Mengubah font semua teks dalam dokumen menjadi Arial ukuran 11."""
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

def change_table_fonts(table):
    """Mengubah font semua teks dalam tabel menjadi Arial ukuran 11."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Fungsi untuk mengubah format tanggal menjadi '12 Maret 2023'
def format_tanggal(tanggal_str):
    bulan_inggris_ke_indonesia = {
        "January": "Januari", "February": "Februari", "March": "Maret",
        "April": "April", "May": "Mei", "June": "Juni",
        "July": "Juli", "August": "Agustus", "September": "September",
        "October": "Oktober", "November": "November", "December": "Desember"
    }
    try:
        dt = datetime.strptime(tanggal_str, '%Y-%m-%d')
        tanggal_format = dt.strftime('%d %B %Y')
        for eng, ind in bulan_inggris_ke_indonesia.items():
            tanggal_format = tanggal_format.replace(eng, ind)
        return tanggal_format
    except ValueError:
        return tanggal_str

# Fungsi untuk mengubah periode menjadi format sesuai aturan
def format_periode(periode_awal, periode_akhir):
    try:
        start = datetime.strptime(periode_awal, '%Y-%m-%d')
        end = datetime.strptime(periode_akhir, '%Y-%m-%d')
        
        # Jika tahun sama dan bulan sama, tampilkan tanggal dan tahun
        if start.year == end.year and start.month == end.month:
            return f"{start.day} s.d {end.day} {format_tanggal(start.strftime('%Y-%m-%d')).split()[-2]} {start.year}"
        
        # Jika tahun sama tapi bulan berbeda
        elif start.year == end.year:
            return f"{format_tanggal(periode_awal).split()[0]} {format_tanggal(periode_awal).split()[1]} s.d {end.day} {format_tanggal(periode_akhir).split()[-2]} {end.year}"
        
        # Jika tahun berbeda
        else:
            return f"{format_tanggal(periode_awal)} s.d {format_tanggal(periode_akhir)}"
    
    except ValueError:
        return f"{periode_awal} s.d {periode_akhir}"

# Fungsi validasi tanggal kegiatan dalam rentang periode penugasan
def validate_tanggal_kegiatan(tanggal_kegiatan, periode_awal, periode_akhir):
    try:
        tanggal = datetime.strptime(tanggal_kegiatan, '%Y-%m-%d')
        start_date = datetime.strptime(periode_awal, '%Y-%m-%d')
        end_date = datetime.strptime(periode_akhir, '%Y-%m-%d')
        return start_date <= tanggal <= end_date
    except ValueError:
        return False  

# Fungsi untuk membuat tabel tanda tangan
def create_signature_table(doc, bagian_ketua_tim, ketua_tim, nama_petugas, lokasi, tanggal_surat):
    # Jika lokasi dan tanggal disediakan, tambahkan paragraf di atas tabel
    if lokasi and tanggal_surat:
    # Menggabungkan lokasi dan tanggal dalam satu paragraf
        paragraph = doc.add_paragraph(f"{lokasi},{format_tanggal(tanggal_surat)}")
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        change_font(doc)

    #Membuat tabel tanda tangan dengan nama Ketua Tim sesuai bagian yang dipilih.
    
    table_paragraph = doc.add_paragraph()
    table_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=2, cols=2)
    table_paragraph._p.addnext(table._tbl)

    # Baris pertama (judul kolom)
    table.cell(0, 0).text = f"Mengetahui,\n{bagian_ketua_tim}\n\n\n"
    table.cell(0, 1).text = "Petugas,"

    # Baris kedua (nama Ketua Tim dan Petugas)
    table.cell(1, 0).text = ketua_tim
    table.cell(1, 1).text = nama_petugas
    change_table_fonts(table)
    # Format teks dan perataan
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.font.size = Pt(10)


    tbl = table._tbl
    for row in tbl.iter_tcs():
        tc_pr = row.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            borders.append(border)
        tc_pr.append(borders)

# Fungsi untuk menghapus baris kosong di tabel Word
def remove_empty_rows(doc):
    for table in doc.tables:
        rows_to_remove = [row for row in table.rows if all(cell.text.strip() == "" for cell in row.cells)]
        for row in rows_to_remove:
            table._tbl.remove(row._tr)

# Fungsi untuk menghapus baris kosong di tabel Word
def remove_empty_rows(doc):
    for table in doc.tables:
        rows_to_remove = [row for row in table.rows if all(cell.text.strip() == "" for cell in row.cells)]
        for row in rows_to_remove:
            table._tbl.remove(row._tr)

@app.route('/')
def home():
    session.setdefault("gambar_count", 1)
    return render_template("form_lapangan.html", gambar_count=session["gambar_count"])

@app.route('/generate', methods=['POST'])
def generate():
    bagian_ketua_tim = request.form.get('bagian_ketua_tim', '')
    ketua_tim = request.form.get('ketua_tim', '')
    petugas = request.form.get('petugas', '')
    tanggal_laporan = request.form.get('tanggal_laporan', '')
    lokasi = request.form.get('lokasi', '')
    periode_awal = (request.form.get('periode_awal') or '').strip()
    periode_akhir = (request.form.get('periode_akhir') or '').strip()
    nomor_surat = request.form['nomor_surat']
    tanggal_surat = request.form['tanggal_surat']
    nama_petugas = request.form['nama_petugas'].strip().replace(" ", "_")
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    template_path = os.path.join(TEMPLATE_FOLDER, 'template.docx')
    
    if not os.path.exists(template_path):
        return "Template tidak ditemukan."
    
    doc = Document(template_path)
    tanggal_kegiatan_list = request.form.getlist('tanggal_kegiatan[]')
    uraian_kegiatan_list = request.form.getlist('uraian_kegiatan[]')
    permasalahan_list = request.form.getlist('permasalahan[]')
    pemecahan_masalah_list = request.form.getlist('pemecahan_masalah[]')
    keterangan_list = request.form.getlist('keterangan[]')

    create_signature_table(doc,bagian_ketua_tim=bagian_ketua_tim, ketua_tim=ketua_tim, nama_petugas=nama_petugas, lokasi=lokasi, tanggal_surat=tanggal_surat)
    doc.add_page_break()
    # Periksa apakah periode_awal dan periode_akhir tidak kosong
    if not periode_awal or not periode_akhir:
        return "Periode penugasan tidak boleh kosong.", 400

    # Tambahkan gambar untuk setiap kegiatan
    for i in range(len(tanggal_kegiatan_list)):
        gambar_list = request.files.getlist(f'gambar_kegiatan_{i}[]')  # Dapatkan semua gambar untuk kegiatan ini
        if gambar_list:
            doc.add_paragraph(f"Dokumentasi untuk Kegiatan pada Tanggal {format_tanggal(tanggal_kegiatan_list[i])}:").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            change_font(doc)
            for gambar in gambar_list:
                if gambar.filename != '':
                    gambar_bytes = gambar.read()  # Membaca gambar sebagai byte stream
                    image_stream = io.BytesIO(gambar_bytes)  # Mengonversi byte stream ke objek file-like
                    
                    # Menambahkan gambar ke dalam dokumen
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Ratakan ke tengah setiap gambar
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(5), height=Inches(4))  # Ukuran gambar disesuaikan
    # Validasi semua tanggal kegiatan (Pastikan ada valid_tanggal_kegiatan)
    valid_tanggal_kegiatan = [t for t in tanggal_kegiatan_list if validate_tanggal_kegiatan(t, periode_awal, periode_akhir)]
    invalid_tanggal_kegiatan = [t for t in tanggal_kegiatan_list if not validate_tanggal_kegiatan(t, periode_awal, periode_akhir)]

    # Debugging: Cek apakah daftar tanggal dikirimkan
    print("Daftar Tanggal Kegiatan yang Diterima:", tanggal_kegiatan_list)
    print("Valid Tanggal Kegiatan:", valid_tanggal_kegiatan)
    print("Invalid Tanggal Kegiatan:", invalid_tanggal_kegiatan)

    if not valid_tanggal_kegiatan:
        return "Tidak ada tanggal kegiatan yang valid dalam periode penugasan.", 400

    # Cek apakah jumlah uraian kegiatan sesuai dengan jumlah tanggal
    if len(valid_tanggal_kegiatan) != len(uraian_kegiatan_list):
        return "Jumlah uraian kegiatan tidak sesuai dengan jumlah tanggal kegiatan.", 400

    for table in doc.tables:
        if 'Nomor Surat Tugas' in table.rows[0].cells[0].text:
            for row in table.rows:
                if 'Nomor Surat Tugas' in row.cells[0].text:
                    row.cells[1].text = nomor_surat
                elif 'Tanggal Surat Tugas' in row.cells[0].text:
                    row.cells[1].text = format_tanggal(tanggal_surat)
                elif 'Nama Petugas' in row.cells[0].text:
                    row.cells[1].text = nama_petugas
                elif 'Periode Penugasan' in row.cells[0].text:
                    row.cells[1].text = format_periode(periode_awal, periode_akhir)
                elif 'Lokasi' in row.cells[0].text:
                    row.cells[1].text = lokasi
            change_table_fonts(table)

    # Loop untuk menambahkan data ke tabel Word
    for table in doc.tables:
        if "No." in table.rows[0].cells[0].text:
            for i, tanggal in enumerate(valid_tanggal_kegiatan):
                row = table.add_row()
                row.cells[0].text = str(i + 1)
                row.cells[1].text = format_tanggal(tanggal)
                row.cells[2].text = uraian_kegiatan_list[i]
                row.cells[3].text = permasalahan_list[i]
                row.cells[4].text = pemecahan_masalah_list[i]
                row.cells[5].text = keterangan_list[i]
            change_table_fonts(table)
            break
    change_font(doc)
    remove_empty_rows(doc)
    # Simpan dokumen ke dalam memori (BytesIO)
    word_mem = io.BytesIO()
    doc.save(word_mem)
    word_mem.seek(0)

    # Tentukan nama file
    os.makedirs('generated_laporan', exist_ok=True)
    file_name = f"laporan pendataan lapangan_{nama_petugas}_{timestamp}.docx"
    file_path = os.path.join("static","generated_laporan", file_name)
    doc.save(file_path)
    # Simpan file secara fisik ke dalam folder
    with open(file_path, 'wb') as f:
        f.write(word_mem.read())
    
    # Simpan path file ke dalam database
    simpan_ke_database(
        nama_file=file_name,
        jenis_form="Form Perjalanan Dinas",  # atau sesuai jenis form
        nama_user=nama_petugas,  # Gantilah dengan nama pengguna yang sesuai
        path_file=file_path
    )

    # Kirim file Word dari memori sebagai lampiran
    return send_file(word_mem, as_attachment=True, download_name=file_name, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)